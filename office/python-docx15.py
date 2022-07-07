from docx import Document
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE

doc = Document()
for i in range(10):
    p = doc.add_paragraph(u'段落 %d' % i)
    style = doc.styles.add_style('UserStyle%d' % i, WD_STYLE_TYPE.PARAGRAPH)
    style.paragraph_format.left_indent = Cm(i)
    p.style = style
    if i == 7:
        style.hidden = False
        style.quick_style = True

for style in doc.styles:
    print(style.name, style.builtin)

doc.paragraphs[3].style = doc.styles['Subtitle']
doc.save('style-4.docx')
