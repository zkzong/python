from docx import Document
from PIL import Image, ImageDraw
from io import BytesIO

document = Document()  # 新建文档
p = document.add_paragraph()  # 添加一个段落
r = p.add_run()  # 添加一个游程
img_size = 20
for x in range(20):
    im = Image.new("RGB", (img_size, img_size), "white")
    draw_obj = ImageDraw.Draw(im)
    draw_obj.ellipse((0, 0, img_size - 1, img_size - 1), fill=255 - x)  # 画圆
    fake_buf_file = BytesIO()  # 用BytesIO将图片保存在内存里，减少磁盘操作
    im.save(fake_buf_file, "png")
    r.add_picture(fake_buf_file)  # 在当前游程中插入图片
    fake_buf_file.close()
    document.save("../file/third/demo.docx")
