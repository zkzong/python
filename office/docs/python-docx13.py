from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

doc = Document()
p = doc.add_paragraph()
text_str = u'好好学习Python，努力做到开发专家，成为最牛都程序员。'
for i, ch in enumerate(text_str):
    run = p.add_run(ch)
    font = run.font
    font.name = u'微软雅黑'
    # bug of python-docx
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    font.bold = (i % 2 == 0)
    font.italic = (i % 3 == 0)
    color = font.color
    color.rgb = RGBColor(i * 10 % 200 + 55, i * 20 % 200 + 55, i * 30 % 200 + 55)

doc.save('style-2.docx')
