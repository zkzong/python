from docx import Document

doc = Document()
doc.add_paragraph(u'Python为什么这么受欢迎？','Title')
doc.add_paragraph(u'作者','Subtitle')
doc.add_paragraph(u'摘要：本文阐明了Python的优势...','Body Text 2')
doc.add_paragraph(u'简单','Heading 1')
doc.add_paragraph(u'易学')
doc.add_paragraph(u'易用','Heading 2')
doc.add_paragraph(u'功能强')
p = doc.add_paragraph(u'贴合小年轻')
p.style = 'Heading 2'
doc.save('demo.docx')