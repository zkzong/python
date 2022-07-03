from docx import Document

document = Document()
table = document.add_table(rows=9,cols=10,style = 'Table Grid')
cell_1 = table.cell(1,2)
cell_2 = table.cell(4,6)
cell_1.merge(cell_2)
document.save('table-1.docx')

document = Document('table-1.docx')
table = document.tables[0]
for row,obj_row in enumerate(table.rows):
    for col,cell in enumerate(obj_row.cells):
        cell.text = cell.text + "%d,%d " % (row,col)

document.save('table-2.docx')